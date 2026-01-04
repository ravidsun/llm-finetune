"""
LangChain-powered data pipeline for document processing and dataset preparation.

Features:
- Multi-format document loading (PDF, DOCX, JSON, JSONL)
- LangChain text splitters for intelligent chunking
- Optional LLM-based QA generation (Claude Haiku)
- Prompt templating with LangChain
"""

import json
import logging
import os
import re
from pathlib import Path
from typing import Any, Dict, Generator, List, Optional, Tuple, Union

from datasets import Dataset, DatasetDict

# LangChain imports
from langchain_core.documents import Document
from langchain_text_splitters import (
    RecursiveCharacterTextSplitter,
    CharacterTextSplitter,
    TokenTextSplitter,
)

from finetune_project.config import Config, DataConfig
from finetune_project.prompt_templates import PromptTemplateManager, create_chat_prompt_from_example
from finetune_project.augmentation import DataAugmenter

logger = logging.getLogger(__name__)


class LangChainDataPipeline:
    """
    Data pipeline using LangChain for document processing.
    
    Handles:
    - PDF, DOCX, and text file loading via LangChain loaders
    - JSON/JSONL loading with schema auto-detection
    - Text chunking using LangChain splitters
    - Optional QA generation using Claude Haiku
    - Data augmentation and formatting
    """

    def __init__(self, config: Config):
        """
        Initialize LangChain data pipeline.
        
        Args:
            config: Full configuration object.
        """
        self.config = config
        self.data_config = config.data
        self.langchain_config = config.data.langchain
        self.prompt_manager = PromptTemplateManager(
            template_dir=self.langchain_config.prompt_template_dir,
            default_system_message=self.langchain_config.default_system_message,
        )
        self._setup_logging()
        self._setup_text_splitter()

    def _setup_logging(self):
        """Configure logging for the pipeline."""
        logging.basicConfig(
            level=logging.INFO,
            format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
        )

    def _setup_text_splitter(self):
        """Initialize LangChain text splitter based on config."""
        pdf_config = self.data_config.pdf
        splitter_type = pdf_config.splitter_type
        
        if splitter_type == "recursive":
            separators = pdf_config.separators or ["\n\n", "\n", ". ", " ", ""]
            self.text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=pdf_config.chunk_size,
                chunk_overlap=pdf_config.chunk_overlap,
                separators=separators,
                length_function=len,
            )
        elif splitter_type == "character":
            self.text_splitter = CharacterTextSplitter(
                chunk_size=pdf_config.chunk_size,
                chunk_overlap=pdf_config.chunk_overlap,
                separator="\n\n",
            )
        elif splitter_type == "token":
            self.text_splitter = TokenTextSplitter(
                chunk_size=pdf_config.chunk_size,
                chunk_overlap=pdf_config.chunk_overlap,
            )
        else:
            # Default to recursive
            self.text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=pdf_config.chunk_size,
                chunk_overlap=pdf_config.chunk_overlap,
            )
        
        logger.info(f"Initialized {splitter_type} text splitter (chunk_size={pdf_config.chunk_size})")

    def prepare(self) -> DatasetDict:
        """
        Main entry point: detect input type and process accordingly.
        
        Returns:
            DatasetDict with 'train' and optionally 'eval' splits.
        """
        input_path = Path(self.data_config.input_path)
        input_type = self._detect_input_type(input_path)
        
        logger.info(f"Detected input type: {input_type}")
        logger.info(f"Processing data from: {input_path}")
        
        if input_type == "json":
            examples = self._process_json(input_path)
        elif input_type in ["pdf", "docx"]:
            examples = self._process_documents(input_path, input_type)
        else:
            raise ValueError(f"Unsupported input type: {input_type}")
        
        # Apply augmentation if enabled
        if self.data_config.augmentation.enabled:
            augmenter = DataAugmenter(
                self.data_config.augmentation,
                seed=self.config.training.seed,
            )
            examples = augmenter.augment_dataset(examples)
        
        # Convert to dataset
        dataset = Dataset.from_list(examples)
        logger.info(f"Created dataset with {len(dataset)} examples")
        
        # Split into train/eval
        dataset_dict = self._split_dataset(dataset)
        
        # Save processed data
        self._save_dataset(dataset_dict)
        
        return dataset_dict

    def _detect_input_type(self, path: Path) -> str:
        """
        Detect input type based on files.

        Args:
            path: Input file or directory path.

        Returns:
            Detected type: 'json', 'pdf', or 'docx'.
        """
        # Convert to string to avoid Path object comparison issues
        input_type_config = str(self.data_config.input_type) if self.data_config.input_type else "auto"
        if input_type_config != "auto":
            return input_type_config
        
        if path.is_file():
            suffix = path.suffix.lower()
            if suffix in ['.json', '.jsonl']:
                return "json"
            elif suffix == '.pdf':
                return "pdf"
            elif suffix in ['.docx', '.doc']:
                return "docx"
            elif suffix in ['.txt', '.md']:
                return "text"
        elif path.is_dir():
            # Check file extensions in directory
            json_files = list(path.glob("*.json")) + list(path.glob("*.jsonl"))
            pdf_files = list(path.glob("*.pdf"))
            docx_files = list(path.glob("*.docx"))
            
            if json_files and not (pdf_files or docx_files):
                return "json"
            elif pdf_files and not json_files:
                return "pdf"
            elif docx_files and not json_files:
                return "docx"
            elif pdf_files or docx_files:
                return "pdf"  # Mixed documents, process as PDF
        
        raise ValueError(f"Cannot detect input type for: {path}")

    # =========================================================================
    # JSON Processing
    # =========================================================================

    def _process_json(self, path: Path) -> List[Dict[str, Any]]:
        """
        Process JSON/JSONL files into training examples.
        
        Args:
            path: Path to JSON file or directory.
            
        Returns:
            List of processed examples.
        """
        raw_examples = self._load_json_files(path)
        schema = self._detect_json_schema(raw_examples)
        logger.info(f"Detected JSON schema: {schema}")
        
        processed = []
        for example in raw_examples:
            formatted = self._format_json_example(example, schema)
            if formatted:
                processed.append(formatted)
        
        logger.info(f"Processed {len(processed)} examples from JSON")
        return processed

    def _load_json_files(self, path: Path) -> List[Dict[str, Any]]:
        """Load JSON/JSONL from file or directory."""
        examples = []
        
        if path.is_file():
            files = [path]
        else:
            files = list(path.glob("*.json")) + list(path.glob("*.jsonl"))
        
        for file_path in files:
            logger.info(f"Loading: {file_path}")
            with open(file_path, 'r', encoding='utf-8') as f:
                if file_path.suffix == '.jsonl':
                    for line in f:
                        line = line.strip()
                        if line:
                            examples.append(json.loads(line))
                else:
                    data = json.load(f)
                    if isinstance(data, list):
                        examples.extend(data)
                    else:
                        examples.append(data)
        
        return examples

    def _detect_json_schema(self, examples: List[Dict[str, Any]]) -> str:
        """Auto-detect JSON schema from examples."""
        if self.data_config.json.schema != "auto":
            return self.data_config.json.schema
        
        if not examples:
            raise ValueError("No examples to detect schema from")
        
        sample = examples[0]
        keys = set(sample.keys())
        
        if 'messages' in keys:
            messages = sample.get('messages', [])
            if messages and isinstance(messages, list):
                if all(isinstance(m, dict) and 'role' in m and 'content' in m for m in messages):
                    return "chatml"
        
        if {'instruction', 'output'}.issubset(keys):
            return "alpaca"
        
        if {'prompt', 'completion'}.issubset(keys):
            return "prompt_completion"
        
        if 'text' in keys:
            return "text_only"
        
        if {'question', 'answer'}.issubset(keys):
            return "qa"
        
        return "custom"

    def _format_json_example(
        self,
        example: Dict[str, Any],
        schema: str,
    ) -> Optional[Dict[str, Any]]:
        """Format a JSON example into training format."""
        json_config = self.data_config.json
        
        try:
            if schema == "alpaca":
                instruction = example.get(json_config.instruction_field, "")
                input_text = example.get(json_config.input_field, "")
                output = example.get(json_config.output_field, "")
                
                if not instruction or not output:
                    return None
                
                user_content = instruction
                if input_text:
                    user_content = f"{instruction}\n\n{input_text}"
                
                return {
                    "messages": [
                        {"role": "user", "content": user_content},
                        {"role": "assistant", "content": output}
                    ]
                }
            
            elif schema == "chatml":
                messages = example.get(json_config.messages_field, [])
                if not messages:
                    return None
                return {"messages": messages}
            
            elif schema == "prompt_completion":
                prompt = example.get(json_config.prompt_field, "")
                completion = example.get(json_config.completion_field, "")
                
                if not prompt or not completion:
                    return None
                
                return {
                    "messages": [
                        {"role": "user", "content": prompt},
                        {"role": "assistant", "content": completion}
                    ]
                }
            
            elif schema == "qa":
                question = example.get("question", "")
                answer = example.get("answer", "")
                
                if not question or not answer:
                    return None
                
                return {
                    "messages": [
                        {"role": "user", "content": question},
                        {"role": "assistant", "content": answer}
                    ]
                }
            
            elif schema == "text_only":
                text = example.get(json_config.text_field, "")
                if not text:
                    return None
                return {"text": text}
            
            elif schema == "custom":
                # Try multiple field mappings
                instruction = example.get(json_config.instruction_field, "")
                input_text = example.get(json_config.input_field, "")
                output = example.get(json_config.output_field, "")
                
                if output:
                    user_content = instruction
                    if input_text:
                        user_content = f"{instruction}\n\n{input_text}"
                    return {
                        "messages": [
                            {"role": "user", "content": user_content},
                            {"role": "assistant", "content": output}
                        ]
                    }
                
                text = example.get(json_config.text_field, "")
                if text:
                    return {"text": text}
                
                return None
            
            else:
                return None
                
        except Exception as e:
            logger.warning(f"Error formatting example: {e}")
            return None

    # =========================================================================
    # Document Processing (PDF, DOCX) using LangChain
    # =========================================================================

    def _process_documents(
        self,
        path: Path,
        doc_type: str,
    ) -> List[Dict[str, Any]]:
        """
        Process documents using LangChain loaders.
        
        For documents without paired labels, creates text-only examples
        for causal LM (continued pretraining).
        
        Optionally generates Q&A pairs using LLM if enabled.
        
        Args:
            path: Path to document(s).
            doc_type: Document type ('pdf' or 'docx').
            
        Returns:
            List of training examples.
        """
        # Load documents using LangChain
        documents = self._load_documents_langchain(path, doc_type)
        logger.info(f"Loaded {len(documents)} documents")
        
        # Split into chunks
        chunks = self.text_splitter.split_documents(documents)
        logger.info(f"Split into {len(chunks)} chunks")
        
        # Process chunks into training examples
        examples = []
        
        # Check if QA generation is enabled
        if self.langchain_config.qa_generation_enabled:
            logger.info("QA generation enabled - generating Q&A pairs from chunks")
            qa_pairs = self._generate_qa_pairs(chunks)
            examples.extend(qa_pairs)
        
        # Also create text-only examples for causal LM
        for chunk in chunks:
            text = chunk.page_content.strip()
            if len(text) >= self.data_config.pdf.min_chunk_size:
                examples.append({"text": text})
        
        logger.info(f"Created {len(examples)} examples from documents")
        return examples

    def _load_documents_langchain(
        self,
        path: Path,
        doc_type: str,
    ) -> List[Document]:
        """
        Load documents using LangChain document loaders.
        
        Args:
            path: Path to document(s).
            doc_type: Document type.
            
        Returns:
            List of LangChain Document objects.
        """
        documents = []
        
        if path.is_file():
            files = [path]
        else:
            if doc_type == "pdf":
                files = list(path.glob("*.pdf"))
            elif doc_type == "docx":
                files = list(path.glob("*.docx")) + list(path.glob("*.doc"))
            else:
                files = list(path.glob("*.*"))
        
        for file_path in files:
            try:
                docs = self._load_single_document(file_path)
                documents.extend(docs)
                logger.info(f"Loaded {len(docs)} pages from {file_path.name}")
            except Exception as e:
                logger.warning(f"Failed to load {file_path}: {e}")
        
        return documents

    def _load_single_document(self, file_path: Path) -> List[Document]:
        """Load a single document file."""
        suffix = file_path.suffix.lower()
        backend = self.data_config.pdf.extraction_backend
        
        if suffix == '.pdf':
            if backend == "pymupdf":
                return self._load_pdf_pymupdf(file_path)
            elif backend == "pdfplumber":
                return self._load_pdf_pdfplumber(file_path)
            elif backend == "unstructured":
                return self._load_with_unstructured(file_path)
            else:
                return self._load_pdf_pymupdf(file_path)
        
        elif suffix in ['.docx', '.doc']:
            return self._load_docx(file_path)
        
        elif suffix in ['.txt', '.md']:
            return self._load_text_file(file_path)
        
        else:
            raise ValueError(f"Unsupported file type: {suffix}")

    def _load_pdf_pymupdf(self, file_path: Path) -> List[Document]:
        """Load PDF using PyMuPDF."""
        try:
            from langchain_community.document_loaders import PyMuPDFLoader
            loader = PyMuPDFLoader(str(file_path))
            return loader.load()
        except ImportError:
            # Fallback to direct PyMuPDF
            import fitz
            doc = fitz.open(file_path)
            documents = []
            
            for page_num, page in enumerate(doc):
                text = page.get_text()
                if text.strip():
                    documents.append(Document(
                        page_content=text,
                        metadata={
                            "source": str(file_path),
                            "page": page_num,
                        }
                    ))
            
            doc.close()
            return documents

    def _load_pdf_pdfplumber(self, file_path: Path) -> List[Document]:
        """Load PDF using pdfplumber."""
        import pdfplumber
        
        documents = []
        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages):
                text = page.extract_text()
                if text and text.strip():
                    documents.append(Document(
                        page_content=text,
                        metadata={
                            "source": str(file_path),
                            "page": page_num,
                        }
                    ))
        
        return documents

    def _load_with_unstructured(self, file_path: Path) -> List[Document]:
        """Load document using Unstructured library."""
        try:
            from langchain_community.document_loaders import UnstructuredFileLoader
            loader = UnstructuredFileLoader(str(file_path))
            return loader.load()
        except ImportError:
            logger.warning("Unstructured not available, falling back to PyMuPDF")
            return self._load_pdf_pymupdf(file_path)

    def _load_docx(self, file_path: Path) -> List[Document]:
        """Load DOCX file."""
        try:
            from langchain_community.document_loaders import Docx2txtLoader
            loader = Docx2txtLoader(str(file_path))
            return loader.load()
        except ImportError:
            # Fallback to python-docx
            from docx import Document as DocxDocument
            doc = DocxDocument(file_path)
            text = "\n\n".join([para.text for para in doc.paragraphs if para.text.strip()])
            return [Document(
                page_content=text,
                metadata={"source": str(file_path)}
            )]

    def _load_text_file(self, file_path: Path) -> List[Document]:
        """Load text or markdown file."""
        with open(file_path, 'r', encoding='utf-8') as f:
            text = f.read()
        
        return [Document(
            page_content=text,
            metadata={"source": str(file_path)}
        )]

    # =========================================================================
    # QA Generation using LLM (Optional)
    # =========================================================================

    def _generate_qa_pairs(
        self,
        chunks: List[Document],
    ) -> List[Dict[str, Any]]:
        """
        Generate Q&A pairs from document chunks using LLM.
        
        This is optional and requires an API key.
        Uses Claude Haiku by default for cost-effectiveness.
        
        Args:
            chunks: List of document chunks.
            
        Returns:
            List of Q&A training examples.
        """
        if not self.langchain_config.qa_generation_enabled:
            return []
        
        # Get LLM
        llm = self._get_qa_llm()
        if llm is None:
            logger.warning("QA LLM not available, skipping QA generation")
            return []
        
        qa_examples = []
        num_pairs = self.langchain_config.qa_pairs_per_chunk
        
        for i, chunk in enumerate(chunks):
            text = chunk.page_content.strip()
            
            # Skip very short chunks
            if len(text) < 200:
                continue
            
            try:
                # Generate Q&A pairs
                pairs = self._generate_qa_for_chunk(llm, text, num_pairs)
                
                for pair in pairs:
                    qa_examples.append({
                        "messages": [
                            {"role": "user", "content": pair["question"]},
                            {"role": "assistant", "content": pair["answer"]}
                        ]
                    })
                
                if (i + 1) % 10 == 0:
                    logger.info(f"Generated QA for {i + 1}/{len(chunks)} chunks")
                    
            except Exception as e:
                logger.warning(f"Failed to generate QA for chunk {i}: {e}")
                continue
        
        logger.info(f"Generated {len(qa_examples)} Q&A pairs from {len(chunks)} chunks")
        return qa_examples

    def _get_qa_llm(self):
        """Get LLM for QA generation."""
        provider = self.langchain_config.qa_llm_provider
        model_name = self.langchain_config.qa_model_name
        
        if provider == "anthropic":
            api_key = os.getenv("ANTHROPIC_API_KEY")
            if not api_key:
                logger.warning("ANTHROPIC_API_KEY not set")
                return None
            
            try:
                from langchain_anthropic import ChatAnthropic
                return ChatAnthropic(
                    model=model_name,
                    temperature=self.langchain_config.qa_temperature,
                    max_tokens=self.langchain_config.qa_max_tokens,
                    api_key=api_key,
                )
            except ImportError:
                logger.warning("langchain-anthropic not installed. Install with: pip install langchain-anthropic")
                return None
        
        elif provider == "openai":
            api_key = os.getenv("OPENAI_API_KEY")
            if not api_key:
                logger.warning("OPENAI_API_KEY not set")
                return None
            
            try:
                from langchain_openai import ChatOpenAI
                return ChatOpenAI(
                    model=model_name or "gpt-4o-mini",
                    temperature=self.langchain_config.qa_temperature,
                    max_tokens=self.langchain_config.qa_max_tokens,
                    api_key=api_key,
                )
            except ImportError:
                logger.warning("langchain-openai not installed. Install with: pip install langchain-openai")
                return None
        
        else:
            logger.warning(f"Unsupported QA LLM provider: {provider}")
            return None

    def _generate_qa_for_chunk(
        self,
        llm,
        text: str,
        num_pairs: int,
    ) -> List[Dict[str, str]]:
        """Generate Q&A pairs for a single chunk."""
        from langchain_core.messages import HumanMessage, SystemMessage
        
        system_message = f"""You are an expert at creating high-quality question-answer pairs from text.
Generate {num_pairs} diverse, informative question-answer pairs based on the given text.

Guidelines:
- Questions should be specific and answerable from the text
- Answers should be accurate and based only on the provided text
- Include a mix of factual, conceptual, and application questions
- Keep answers concise but complete

Output ONLY a valid JSON array with objects containing "question" and "answer" fields. No other text."""

        human_message = f"""Generate {num_pairs} question-answer pairs from this text:

---
{text}
---

Return ONLY a valid JSON array."""

        messages = [
            SystemMessage(content=system_message),
            HumanMessage(content=human_message),
        ]
        
        response = llm.invoke(messages)
        content = response.content.strip()
        
        # Parse JSON response
        # Handle potential markdown code blocks
        if content.startswith("```"):
            content = re.sub(r'^```(?:json)?\n?', '', content)
            content = re.sub(r'\n?```$', '', content)
        
        pairs = json.loads(content)
        
        # Validate format
        if not isinstance(pairs, list):
            pairs = [pairs]
        
        valid_pairs = []
        for pair in pairs:
            if isinstance(pair, dict) and "question" in pair and "answer" in pair:
                valid_pairs.append({
                    "question": str(pair["question"]),
                    "answer": str(pair["answer"]),
                })
        
        return valid_pairs

    # =========================================================================
    # Dataset Operations
    # =========================================================================

    def _split_dataset(self, dataset: Dataset) -> DatasetDict:
        """Split dataset into train and eval sets."""
        if self.data_config.shuffle:
            dataset = dataset.shuffle(seed=self.config.training.seed)
        
        train_ratio = self.data_config.train_split
        
        if train_ratio >= 1.0:
            return DatasetDict({"train": dataset})
        
        split = dataset.train_test_split(
            test_size=1.0 - train_ratio,
            seed=self.config.training.seed
        )
        
        return DatasetDict({
            "train": split["train"],
            "eval": split["test"]
        })

    def _save_dataset(self, dataset_dict: DatasetDict):
        """Save processed dataset to disk."""
        output_path = Path(self.data_config.output_dataset_path)
        output_path.mkdir(parents=True, exist_ok=True)

        # Save as HF dataset (convert Path to string for fsspec compatibility)
        dataset_dict.save_to_disk(str(output_path))
        logger.info(f"Saved dataset to: {output_path}")
        
        # Also save as JSONL for inspection
        for split_name, split_data in dataset_dict.items():
            jsonl_path = output_path / f"{split_name}.jsonl"
            with open(jsonl_path, 'w', encoding='utf-8') as f:
                for example in split_data:
                    f.write(json.dumps(example, ensure_ascii=False) + '\n')
            logger.info(f"Saved {split_name} JSONL: {jsonl_path} ({len(split_data)} examples)")

    def load_prepared_dataset(self) -> DatasetDict:
        """Load previously prepared dataset."""
        output_path = Path(self.data_config.output_dataset_path)

        if not output_path.exists():
            raise FileNotFoundError(
                f"Prepared dataset not found at {output_path}. "
                "Run 'prepare-data' first."
            )

        from datasets import load_from_disk
        # Convert Path to string for fsspec compatibility
        return load_from_disk(str(output_path))
